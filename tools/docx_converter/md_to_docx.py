from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json, re

BASE = Path(__file__).resolve().parent

def set_east_asia_font(style_or_run, font_name="Noto Sans CJK JP"):
    rpr = None
    if hasattr(style_or_run, "_element"):
        el = style_or_run._element
        rpr = el.get_or_add_rPr() if hasattr(el, "get_or_add_rPr") else None
    if rpr is not None:
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def build():
    md_text = (BASE / "00_full_document.md").read_text(encoding="utf-8")
    sva = json.loads((BASE / "ipower-engine.sva.instance.json").read_text(encoding="utf-8"))
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    set_east_asia_font(normal, "Noto Sans CJK JP")

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x1D, 0x20, 0x87)),
        ("Heading 1", 15.5, RGBColor(0x00, 0x68, 0xB6)),
        ("Heading 2", 12.2, RGBColor(0x1D, 0x20, 0x87)),
        ("Heading 3", 10.8, RGBColor(0x00, 0x68, 0xB6)),
    ]:
        st = styles[style_name]
        st.font.name = "Aptos"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        set_east_asia_font(st, "Noto Sans CJK JP")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("Page ").font.size = Pt(9)
    add_page_number(footer_p)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HEMS SVA設計書")
    r.font.size = Pt(24)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sva["document_info"]["project_name"])
    r2.font.size = Pt(14)
    r2.bold = True

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("文書名", "HEMS SVA設計書"),
        ("プロジェクト", sva["document_info"]["project_name"]),
        ("バージョン", sva["document_info"]["version"]),
        ("更新日", sva["document_info"].get("last_updated", "")),
    ]):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v

    doc.add_paragraph("")
    desc = doc.add_paragraph()
    desc.add_run("概要: ").bold = True
    desc.add_run("JSON Schema / instance JSON / Markdown を原本にし、客先提出向け DOCX を生成するための変換成果物。")
    doc.add_page_break()

    blank_count = 0
    for line in md_text.splitlines():
        stripped = line.rstrip()
        if not stripped:
            blank_count += 1
            if blank_count <= 1:
                doc.add_paragraph("")
            continue
        blank_count = 0
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    out_path = BASE / "i-power-engine_hems_sva_design_auto.docx"
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    build()
